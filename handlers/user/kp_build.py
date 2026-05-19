# handlers/user/kp_build.py
from __future__ import annotations

import asyncio
import logging
import os
import re
import time
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Optional
from urllib.parse import urlsplit

import aiohttp
from aiogram import F, Router
from aiogram.exceptions import TelegramEntityTooLarge
from aiogram.filters import BaseFilter
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup
from aiogram.types import (
    CallbackQuery,
    FSInputFile,
    InlineKeyboardButton,
    Message,
)
from aiogram.utils.keyboard import InlineKeyboardBuilder

from callbacks import UserKpCb
from config import Settings
from db import (
    get_user,
    is_admin,
    get_kp_session,
    list_kp_items,
    count_kp_items,
    get_product,   # ✅ нужно для проверки наличия/stock
)
from services.image_store import default_image_path
from services.product_enrich import enrich_from_url  # ✅ подтягиваем данные по ссылке

# PDF (reportlab)
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.utils import ImageReader
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    Image as RLImage,
    KeepTogether,
    PageBreak,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.fonts import addMapping
from xml.sax.saxutils import escape as xml_escape

# EXCEL (openpyxl)
from openpyxl import Workbook
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

# WORD (python-docx)
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from PIL import Image, ImageOps
except Exception:
    Image = None
    ImageOps = None


router = Router()
log = logging.getLogger(__name__)

# -------------------- Branding / Design --------------------

COMPANY_NAME = "Goodwin"
LOGO_DIR = Path("media") / "logo" / "png"

ACCENT_GREEN = "#1F331A"
ACCENT_WINE = "#68162A"

BG_SOFT = "#F8FAFC"
TEXT_DARK = "#0F172A"
TEXT_MUTED = "#475569"
BORDER_SOFT = "#E2E8F0"
HEADER_SOFT = "#F1F5F9"


TELEGRAM_DOCUMENT_LIMIT_BYTES = 48 * 1024 * 1024
KP_IMAGE_MAX_SIDE = 900
KP_IMAGE_QUALITY = 72

# -------------------- Filter --------------------

class IsApprovedUser(BaseFilter):
    async def __call__(self, event: Message | CallbackQuery, settings: Settings) -> bool:
        uid = event.from_user.id
        if await is_admin(settings.db_path, uid, settings.admin_ids):
            return False
        user = await get_user(settings.db_path, uid)
        return bool(user and user.get("status") == "approved")


# -------------------- FSM --------------------

class KpBuildForm(StatesGroup):
    choosing_mode = State()
    waiting_qty = State()
    waiting_url = State()
    review = State()       # ✅ новое: проверка пропусков перед генерацией


# -------------------- Small helpers (clean chat) --------------------

async def safe_delete_message(bot, chat_id: int, message_id: int) -> None:
    if not message_id:
        return
    try:
        await bot.delete_message(chat_id, message_id)
    except Exception:
        pass


async def safe_edit_text(message: Message, text: str, reply_markup=None) -> None:
    try:
        await message.edit_text(text, reply_markup=reply_markup)
    except Exception:
        try:
            await message.delete()
        except Exception:
            pass
        await message.bot.send_message(message.chat.id, text, reply_markup=reply_markup)


# -------------------- Utils --------------------

def _money(v: Any) -> str:
    if v is None or v == "":
        return "—"
    try:
        return f"{float(v):,.2f}".replace(",", " ").replace(".", ",")
    except Exception:
        return str(v)


def _safe_float(v: Any) -> float:
    try:
        if v is None or v == "":
            return 0.0
        return float(v)
    except Exception:
        return 0.0


def _short(s: str | None, n: int = 70) -> str:
    t = (s or "").strip()
    t = re.sub(r"\s+", " ", t)
    return (t[: n - 1] + "…") if len(t) > n else t


def _rl_escape(s: str | None) -> str:
    return xml_escape((s or "").strip(), {"'": "&apos;", '"': "&quot;"})


def _kp_dir() -> str:
    p = Path("media") / "kp"
    p.mkdir(parents=True, exist_ok=True)
    return str(p)


def _kp_img_dir() -> str:
    p = Path("media") / "kp" / "images"
    p.mkdir(parents=True, exist_ok=True)
    return str(p)


def _find_logo_path() -> str | None:
    try:
        if not LOGO_DIR.exists():
            return None
        for p in sorted(LOGO_DIR.glob("*.png")):
            if p.is_file():
                return str(p)
        return None
    except Exception:
        return None


def _build_base_filename() -> str:
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"KP_{COMPANY_NAME}_{ts}"


def _register_russian_fonts() -> tuple[str, str, str]:
    """
    ✅ PDF: Times New Roman (с кириллицей) через TTF.
    Если TimesNewRoman.ttf не найден — fallback на DejaVuSerif (тоже с кириллицей).
    Чтобы на Railway было 100%: положи шрифты в media/fonts/
    """
    local_dir = Path("media") / "fonts"

    candidates_regular = [
        # --- локально (РЕКОМЕНДУЕТСЯ для Railway) ---
        str(local_dir / "TimesNewRoman.ttf"),
        str(local_dir / "Times New Roman.ttf"),
        str(local_dir / "TimesNewRomanPSMT.ttf"),
        str(local_dir / "times.ttf"),

        # --- Linux пути (иногда есть msttcorefonts) ---
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/times.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/TimesNewRoman.ttf",

        # --- fallback serif с кириллицей ---
        str(local_dir / "DejaVuSerif.ttf"),
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
        str(local_dir / "NotoSerif-Regular.ttf"),
        "/usr/share/fonts/truetype/noto/NotoSerif-Regular.ttf",

        # --- Windows ---
        r"C:\Windows\Fonts\times.ttf",
        r"C:\Windows\Fonts\Times.ttf",
    ]

    candidates_bold = [
        # --- локально ---
        str(local_dir / "TimesNewRoman.ttf"),
        str(local_dir / "Times New Roman Bold.ttf"),
        str(local_dir / "TimesNewRomanPS-BoldMT.ttf"),
        str(local_dir / "timesbd.ttf"),

        # --- Linux ---
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/timesbd.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/TimesNewRoman.ttf",

        # --- fallback serif bold ---
        str(local_dir / "DejaVuSerif-Bold.ttf"),
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
        str(local_dir / "NotoSerif-Bold.ttf"),
        "/usr/share/fonts/truetype/noto/NotoSerif-Bold.ttf",

        # --- Windows ---
        r"C:\Windows\Fonts\timesbd.ttf",
        r"C:\Windows\Fonts\Timesbd.ttf",
    ]

    reg_path = next((p for p in candidates_regular if p and os.path.exists(p)), None)
    bold_path = next((p for p in candidates_bold if p and os.path.exists(p)), None)

    if not reg_path:
        # чтобы не падало (но кириллица может стать "квадратами")
        return ("Helvetica", "Helvetica", "Helvetica-Bold")

    family = "KPTimes"
    regular = "KPTimes"
    bold = "KPTimes-Bold"

    try:
        if regular not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(regular, reg_path))

        if bold_path:
            if bold not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(bold, bold_path))
        else:
            bold = regular

        addMapping(family, 0, 0, regular)
        addMapping(family, 1, 0, bold)
        addMapping(family, 0, 1, regular)
        addMapping(family, 1, 1, bold)

        return (family, regular, bold)
    except Exception:
        return ("Helvetica", "Helvetica", "Helvetica-Bold")





def _file_size_mb(path: str) -> float:
    try:
        return round(os.path.getsize(path) / 1024 / 1024, 2)
    except Exception:
        return 0.0


def _compress_image_for_pdf(
    path: str,
    *,
    max_side: int = KP_IMAGE_MAX_SIDE,
    quality: int = KP_IMAGE_QUALITY,
) -> str:
    """
    Сжимает картинку перед вставкой в PDF.
    Важно: визуальное уменьшение в ReportLab не уменьшает вес исходного файла.
    """
    try:
        if Image is None or ImageOps is None:
            return path

        src = Path(path)
        if not src.exists() or not src.is_file():
            return path

        if src.suffix.lower() == ".svg":
            return path

        stat = src.stat()
        cache_dir = Path(_kp_img_dir()) / "compressed"
        cache_dir.mkdir(parents=True, exist_ok=True)

        out_name = f"{src.stem}_{int(stat.st_mtime)}_{stat.st_size}_{max_side}_{quality}.jpg"
        out_path = cache_dir / out_name

        if out_path.exists() and out_path.stat().st_size > 0:
            return str(out_path)

        with Image.open(src) as img:
            img = ImageOps.exif_transpose(img)

            if img.mode in ("RGBA", "LA", "P"):
                bg = Image.new("RGB", img.size, "white")
                if img.mode == "RGBA":
                    bg.paste(img, mask=img.getchannel("A"))
                else:
                    bg.paste(img.convert("RGBA"), mask=img.convert("RGBA").getchannel("A"))
                img = bg
            else:
                img = img.convert("RGB")

            resampling = getattr(getattr(Image, "Resampling", Image), "LANCZOS", 1)
            img.thumbnail((max_side, max_side), resampling)

            img.save(
                out_path,
                "JPEG",
                quality=quality,
                optimize=True,
                progressive=True,
            )

        if out_path.exists() and out_path.stat().st_size > 0:
            if out_path.stat().st_size < stat.st_size:
                return str(out_path)

        return path
    except Exception:
        return path


def _fit_image(path: str, max_w_mm: float, max_h_mm: float) -> RLImage:
    max_w = max_w_mm * mm
    max_h = max_h_mm * mm
    try:
        safe_path = _compress_image_for_pdf(path)
        img_reader = ImageReader(safe_path)
        iw, ih = img_reader.getSize()
        if not iw or not ih:
            raise ValueError("bad image size")
        scale = min(max_w / iw, max_h / ih)
        w = iw * scale
        h = ih * scale
        return RLImage(safe_path, width=w, height=h)
    except Exception:
        fallback = _compress_image_for_pdf(default_image_path())
        return RLImage(fallback, width=max_w, height=max_h)


def _item_title(it: dict[str, Any]) -> str:
    return (it.get("title") or it.get("description") or "—").strip()


def _stock_of_item(it: dict[str, Any]) -> int | None:
    pid = it.get("product_id")
    if not pid:
        return None
    try:
        return int(it.get("_stock_qty"))
    except Exception:
        return None


def _is_empty(v: Any) -> bool:
    if v is None:
        return True
    if isinstance(v, str) and v.strip() in ("", "—"):
        return True
    return False


def _is_valid_url(s: str) -> bool:
    s = (s or "").strip()
    if not s:
        return False
    try:
        u = urlsplit(s)
        return u.scheme in ("http", "https") and bool(u.netloc)
    except Exception:
        return False


async def _download_image(url: str, out_path: str, timeout_s: int = 25) -> bool:
    """
    Скачиваем картинку по URL в out_path.
    Возвращает True если успех.
    """
    try:
        t = aiohttp.ClientTimeout(total=timeout_s)
        async with aiohttp.ClientSession(timeout=t) as session:
            async with session.get(url, allow_redirects=True) as resp:
                if resp.status != 200:
                    return False
                ct = (resp.headers.get("Content-Type") or "").lower()
                if "image" not in ct:
                    # иногда сервер не ставит content-type; всё равно попробуем
                    pass
                data = await resp.read()
                if not data or len(data) < 200:
                    return False
                Path(os.path.dirname(out_path)).mkdir(parents=True, exist_ok=True)
                with open(out_path, "wb") as f:
                    f.write(data)
                return True
    except Exception:
        return False


async def _ensure_item_image_downloaded(it: dict[str, Any], uid: int, idx: int) -> None:
    """
    Если у позиции нет image_path, но есть image_url — пробуем скачать, чтобы PDF мог вставить фото.
    """
    try:
        img_path = (it.get("image_path") or "").strip()
        if img_path and os.path.exists(img_path):
            return

        img_url = (it.get("image_url") or "").strip()
        if not img_url:
            return

        # сохраняем в media/kp/images
        ext = ".jpg"
        # иногда можно извлечь расширение
        try:
            p = urlsplit(img_url).path.lower()
            if p.endswith(".png"):
                ext = ".png"
            elif p.endswith(".webp"):
                ext = ".webp"
            elif p.endswith(".jpeg"):
                ext = ".jpeg"
            elif p.endswith(".jpg"):
                ext = ".jpg"
        except Exception:
            pass

        fname = f"kp_{uid}_{int(time.time())}_{idx}{ext}"
        out_path = os.path.join(_kp_img_dir(), fname)

        ok = await _download_image(img_url, out_path)
        if ok and os.path.exists(out_path):
            it["image_path"] = out_path
    except Exception:
        return


async def _enrich_items_with_stock(settings: Settings, items: list[dict[str, Any]]) -> None:
    """
    Обогащаем позиции КП данными из products по product_id:
      - _stock_qty
      - code (артикул)
      - image_url / image_path (чтобы в КП было фото)
    """
    import os
    import json

    for it in items:
        it["_stock_qty"] = None

        # 0) если это web-позиция (без product_id), попробуем вытащить image/code из extra_json (если там есть)
        try:
            ej = it.get("extra_json")
            if isinstance(ej, str) and ej.strip():
                d = json.loads(ej)
                if isinstance(d, dict):
                    # code
                    if _is_empty(it.get("code")):
                        for k in ("code", "article", "sku", "vendor_code"):
                            v = d.get(k)
                            if isinstance(v, str) and v.strip():
                                it["code"] = v.strip()
                                break
                    # image_url
                    if _is_empty(it.get("image_url")):
                        for k in ("image_url", "image", "img", "photo"):
                            v = d.get(k)
                            if isinstance(v, str) and v.strip():
                                it["image_url"] = v.strip()
                                break
        except Exception:
            pass

        pid = it.get("product_id")
        if not pid:
            continue

        # 1) тянем продукт из БД
        try:
            p = await get_product(settings.db_path, int(pid))
        except Exception:
            p = None

        if not p:
            continue

        # 2) stock
        try:
            x = p.get("stock_qty")
            it["_stock_qty"] = int(x) if x is not None and x != "" else 0
        except Exception:
            it["_stock_qty"] = None

        # 3) code
        try:
            pc = p.get("code")
            if pc and _is_empty(it.get("code")):
                it["code"] = str(pc).strip()
        except Exception:
            pass

        # 4) image_path (если файл реально есть)
        try:
            ip = (p.get("image_path") or "").strip()
            if _is_empty(it.get("image_path")) and ip and os.path.exists(ip):
                it["image_path"] = ip
        except Exception:
            pass

        # 5) image_url (чтобы можно было скачать в момент генерации КП)
        try:
            iu = (p.get("image_url") or "").strip()
            if _is_empty(it.get("image_url")) and iu:
                it["image_url"] = iu
        except Exception:
            pass



# -------------------- Totals --------------------

@dataclass
class KpTotals:
    qty_total: int | None
    sum_price: float | None
    sum_final: float | None

    @property
    def economy(self) -> float | None:
        if self.sum_price is None or self.sum_final is None:
            return None
        e = self.sum_price - self.sum_final
        return e if e > 0 else 0.0


def _calc_totals(items: list[dict[str, Any]]) -> KpTotals:
    qty_sum = 0
    sp = 0.0
    sf = 0.0
    any_qty = False

    for it in items:
        q = it.get("doc_qty")
        if q is None:
            continue
        any_qty = True
        qty_sum += int(q)

        price = _safe_float(it.get("price"))
        final = _safe_float(it.get("final_price") if it.get("final_price") is not None else it.get("price"))

        sp += price * int(q)
        sf += final * int(q)

    if not any_qty:
        return KpTotals(qty_total=None, sum_price=None, sum_final=None)

    return KpTotals(qty_total=qty_sum, sum_price=sp, sum_final=sf)


# -------------------- PDF Drawing --------------------

def _draw_header(canvas, doc, font_regular: str, font_bold: str, meta: dict[str, str], logo_path: str | None) -> None:
    canvas.saveState()
    w, h = A4

    header_h = 24 * mm
    canvas.setFillColor(colors.HexColor(BG_SOFT))
    canvas.rect(0, h - header_h, w, header_h, stroke=0, fill=1)

    stripe_h = 2.2 * mm
    canvas.setFillColor(colors.HexColor(ACCENT_GREEN))
    canvas.rect(0, h - header_h, w * 0.6, stripe_h, stroke=0, fill=1)
    canvas.setFillColor(colors.HexColor(ACCENT_WINE))
    canvas.rect(w * 0.6, h - header_h, w * 0.4, stripe_h, stroke=0, fill=1)

    if logo_path and os.path.exists(logo_path):
        try:
            safe_logo_path = _compress_image_for_pdf(logo_path, max_side=700, quality=78)
            img = ImageReader(safe_logo_path)
            iw, ih = img.getSize()
            target_h = 14 * mm
            scale = target_h / float(ih)
            target_w = float(iw) * scale
            canvas.drawImage(img, 18 * mm, h - 18.5 * mm, width=target_w, height=target_h, mask="auto")
        except Exception:
            pass

    canvas.setFillColor(colors.HexColor(TEXT_DARK))
    canvas.setFont(font_bold, 13)
    canvas.drawRightString(w - 18 * mm, h - 15.5 * mm, "Коммерческое предложение")

    canvas.setFillColor(colors.HexColor(TEXT_MUTED))
    canvas.setFont(font_regular, 9.5)
    canvas.drawRightString(w - 18 * mm, h - 20.0 * mm, meta.get("date", ""))

    canvas.setStrokeColor(colors.HexColor(BORDER_SOFT))
    canvas.setLineWidth(1)
    canvas.line(18 * mm, 16 * mm, w - 18 * mm, 16 * mm)

    canvas.setFillColor(colors.HexColor("#94A3B8"))
    canvas.setFont(font_regular, 8.5)
    canvas.drawString(18 * mm, 11 * mm, meta.get("footer_left", ""))
    canvas.drawRightString(w - 18 * mm, 11 * mm, f"Стр. {doc.page}")

    canvas.restoreState()


def create_kp_pdf(
    *,
    out_path: str,
    user: dict[str, Any] | None,
    items: list[dict[str, Any]],
    logo_path: str | None,
) -> None:
    family, font_regular, font_bold = _register_russian_fonts()

    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "base",
        parent=styles["Normal"],
        fontName=family,
        fontSize=10.8,
        leading=14,
        textColor=colors.HexColor(TEXT_DARK),
    )
    h2 = ParagraphStyle(
        "h2",
        parent=base,
        fontSize=12.4,
        leading=16,
        spaceBefore=8,
        spaceAfter=6,
    )
    small = ParagraphStyle(
        "small",
        parent=base,
        fontSize=9.6,
        leading=12,
        textColor=colors.HexColor(TEXT_MUTED),
    )
    name_cell_style = ParagraphStyle(
        "name_cell",
        parent=base,
        fontSize=9.4,
        leading=11.2,
        wordWrap="CJK",
    )

    doc = SimpleDocTemplate(
        out_path,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=30 * mm,
        bottomMargin=22 * mm,
        title="Коммерческое предложение",
        author=COMPANY_NAME,
    )

    today = datetime.now().strftime("%d.%m.%Y")
    kp_no = datetime.now().strftime("%Y-%m-%d/%H%M")

    client_line = ""
    if user:
        fio = (user.get("full_name") or "").strip()
        phone = (user.get("phone") or "").strip()
        reg_type = (user.get("reg_type") or "").strip()
        parts = [p for p in [fio, reg_type, phone] if p]
        client_line = " • ".join(parts)

    totals = _calc_totals(items)
    meta = {"date": today, "footer_left": f"КП № {kp_no}"}

    story: list[Any] = []

    # ---------- Title ----------
    story.append(
        Paragraph(
            f"<b>{_rl_escape(COMPANY_NAME)}</b>",
            ParagraphStyle("t", parent=base, fontSize=16.5, leading=20),
        )
    )
    story.append(Paragraph("Документ сформирован автоматически в системе.", small))
    story.append(Spacer(1, 10))

    # ---------- Info block ----------
    left_lines = [
        f"<b>КП №:</b> {_rl_escape(kp_no)}",
        f"<b>Дата:</b> {_rl_escape(today)}",
        f"<b>Позиции:</b> {_rl_escape(str(len(items)))}",
    ]
    if client_line:
        left_lines.append(f"<b>Клиент:</b> {_rl_escape(client_line)}")

    if totals.qty_total is None:
        right_lines = [
            "<b>Кол-во (итого):</b> —",
            "<b>Итого со скидкой:</b> —",
            "<b>Итого без скидки:</b> —",
            "<b>Экономия:</b> —",
        ]
        note = Paragraph("💡 Кол-во не задано — итоги не рассчитываются.", small)
    else:
        right_lines = [
            f"<b>Кол-во (итого):</b> {_rl_escape(str(totals.qty_total))}",
            f"<b>Итого со скидкой:</b> {_rl_escape(_money(totals.sum_final))}",
            f"<b>Итого без скидки:</b> {_rl_escape(_money(totals.sum_price))}",
            f"<b>Экономия:</b> {_rl_escape(_money(totals.economy))}",
        ]
        note = None

    info_tbl = Table(
        [[Paragraph("<br/>".join(left_lines), base), Paragraph("<br/>".join(right_lines), base)]],
        colWidths=[88 * mm, 78 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(HEADER_SOFT)),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(BORDER_SOFT)),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 10),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        ),
    )
    story.append(info_tbl)
    if note:
        story.append(Spacer(1, 6))
        story.append(note)
    story.append(Spacer(1, 12))

    # ---------- Conditions ----------
    story.append(Paragraph("Условия", h2))
    cond_tbl = Table(
        [["Срок поставки", "—"], ["Оплата", "—"], ["Доставка", "—"]],
        colWidths=[45 * mm, 121 * mm],
        style=TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.white),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(BORDER_SOFT)),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor(BORDER_SOFT)),
                ("FONTNAME", (0, 0), (-1, -1), family),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor(TEXT_DARK)),
                ("TEXTCOLOR", (1, 0), (1, -1), colors.HexColor(TEXT_MUTED)),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        ),
    )
    story.append(cond_tbl)
    story.append(Spacer(1, 12))

    # ---------- Summary table ----------
    story.append(Paragraph("Состав КП", h2))
    rows = [["№", "Артикул", "Наименование", "Кол-во", "Цена", "Сумма"]]

    max_rows = 12
    for i, it in enumerate(items[:max_rows], start=1):
        qty = it.get("doc_qty")

        code = (
            it.get("code")
            or it.get("article")
            or it.get("sku")
            or it.get("vendor_code")
            or (it.get("product_code") if isinstance(it.get("product_code"), str) else None)
            or "—"
        )
        code = str(code).strip() if code is not None else "—"

        title_s = _short(_item_title(it), 60)
        final = _safe_float(it.get("final_price") if it.get("final_price") is not None else it.get("price"))
        sum_final = None if (qty is None or final == 0.0) else (final * int(qty))

        rows.append([
            str(i),
            _rl_escape(code),
            Paragraph(_rl_escape(title_s), name_cell_style),
            str(qty) if qty is not None else "—",
            _rl_escape(_money(final) if final else "—"),
            _rl_escape(_money(sum_final) if sum_final is not None else "—"),
        ])

    if len(items) > max_rows:
        rows.append(["", "", f"… и ещё {len(items) - max_rows} поз.", "", "", ""])

    summary_tbl = Table(
        rows,
        colWidths=[8 * mm, 22 * mm, 80 * mm, 14 * mm, 24 * mm, 18 * mm],
        repeatRows=1,
        style=TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), family),
                ("FONTSIZE", (0, 0), (-1, 0), 10),
                ("FONTSIZE", (0, 1), (-1, -1), 9.4),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(HEADER_SOFT)),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(BORDER_SOFT)),
                ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor(BORDER_SOFT)),
                ("ALIGN", (0, 0), (0, -1), "CENTER"),
                ("ALIGN", (1, 1), (1, -1), "CENTER"),
                ("ALIGN", (3, 1), (3, -1), "CENTER"),
                ("ALIGN", (4, 1), (-1, -1), "RIGHT"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]
        ),
    )
    story.append(summary_tbl)

    # ---------- Cards pages ----------
    story.append(PageBreak())

    card_style = ParagraphStyle(
        "card_style",
        parent=base,
        fontSize=10.2,
        leading=12.4,
        textColor=colors.HexColor(TEXT_DARK),
    )

    usable_w = (A4[0] - 18 * mm - 18 * mm)
    img_col = 44 * mm
    text_col = usable_w - img_col

    CARD_H_MM = 52
    GAP_MM = 6

    # ✅ ВОТ ЭТУ ВЛОЖЕННУЮ ФУНКЦИЮ И НУЖНО БЫЛО ИСПРАВИТЬ
    def make_card_big(idx: int, it: dict[str, Any]) -> Table:
        title_s = _short(_item_title(it), 160)

        desc = (it.get("description") or "").strip()
        desc = re.sub(r"\n{3,}", "\n\n", desc).strip()
        desc = desc[:320]
        desc_html = _rl_escape(desc).replace("\n", "<br/>") if desc else ""

        url = (it.get("url") or "").strip()
        qty = it.get("doc_qty")
        final = _safe_float(it.get("final_price") if it.get("final_price") is not None else it.get("price"))

        img_path = (it.get("image_path") or "").strip()
        if not img_path or not os.path.exists(img_path):
            img_path = default_image_path()
        img = _fit_image(img_path, max_w_mm=36, max_h_mm=36)

        meta_bits = []
        meta_bits.append(f"Кол-во: <b>{_rl_escape(str(qty))}</b>" if qty is not None else "Кол-во: <b>—</b>")
        if final:
            meta_bits.append(f"Цена: <b>{_rl_escape(_money(final))}</b>")
        if qty is not None and final:
            meta_bits.append(f"Сумма: <b>{_rl_escape(_money(final * int(qty)))}</b>")
        meta_html = (" • ".join(meta_bits)) if meta_bits else ""

        url_html = ""
        if url:
            url_href = _rl_escape(url)
            url_text = _rl_escape(_short(url, 60))
            url_html = (
                "Ссылка:<br/>"
                f"<link href='{url_href}' color='{ACCENT_WINE}'>{url_text}</link>"
            )

        right_html = (
            f"<b>{idx}. {_rl_escape(title_s)}</b>"
            + (f"<br/><br/>{desc_html}" if desc_html else "")
            + (f"<br/><br/>{meta_html}" if meta_html else "")
            + (f"<br/><br/>{url_html}" if url_html else "")
        )
        right = Paragraph(right_html, card_style)

        card = Table(
            [[img, right]],
            colWidths=[img_col, text_col],
            rowHeights=[CARD_H_MM * mm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(BG_SOFT)),
                    ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor(BORDER_SOFT)),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 10),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 10),
                    ("VALIGN", (0, 0), (0, 0), "MIDDLE"),
                    ("VALIGN", (1, 0), (1, 0), "TOP"),
                ]
            ),
        )
        return card

    for base_i in range(0, len(items), 4):
        chunk = items[base_i: base_i + 4]

        page_blocks: list[Any] = []
        for j, it in enumerate(chunk, start=1):
            page_blocks.append(make_card_big(base_i + j, it))
            if j != len(chunk):
                page_blocks.append(Spacer(1, GAP_MM))

        story.append(KeepTogether(page_blocks))

        if base_i + 4 < len(items):
            story.append(PageBreak())

    doc.build(
        story,
        onFirstPage=lambda c, d: _draw_header(c, d, font_regular, font_bold, meta, logo_path),
        onLaterPages=lambda c, d: _draw_header(c, d, font_regular, font_bold, meta, logo_path),
    )


# -------------------- DOCX helpers --------------------

def _docx_set_cell_shading(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex.replace("#", ""))
    tcPr.append(shd)


def _docx_set_cell_borders(cell, bottom_hex: str | None = None) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)

    if bottom_hex:
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "10")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), bottom_hex.replace("#", ""))
        tcBorders.append(bottom)


def _docx_set_default_font(doc: Document, font_name: str = "Times New Roman", font_size_pt: float = 10.5) -> None:
    try:
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size_pt)

        rPr = style.element.rPr
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            style.element.append(rPr)

        rFonts = rPr.rFonts
        rFonts.set(qn("w:ascii"), font_name)
        rFonts.set(qn("w:hAnsi"), font_name)
        rFonts.set(qn("w:cs"), font_name)
        rFonts.set(qn("w:eastAsia"), font_name)
    except Exception:
        pass



def create_kp_docx(
    *,
    out_path: str,
    user: dict[str, Any] | None,
    items: list[dict[str, Any]],
    logo_path: str | None,
) -> None:
    doc = Document()
    _docx_set_default_font(doc, "Arial", 10.5)

    header = doc.add_table(rows=1, cols=2)
    header.autofit = True
    left, right = header.rows[0].cells

    _docx_set_cell_shading(left, BG_SOFT)
    _docx_set_cell_shading(right, BG_SOFT)
    _docx_set_cell_borders(left, ACCENT_GREEN)
    _docx_set_cell_borders(right, ACCENT_WINE)

    if logo_path and os.path.exists(logo_path):
        p = left.paragraphs[0]
        run = p.add_run()
        try:
            run.add_picture(logo_path, width=Mm(35))
        except Exception:
            pass
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Коммерческое предложение")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Arial"

    p2 = right.add_paragraph(datetime.now().strftime("%d.%m.%Y"))
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p2.runs:
        p2.runs[0].font.size = Pt(9)
        p2.runs[0].font.name = "Arial"

    doc.add_paragraph("")

    today = datetime.now().strftime("%d.%m.%Y")
    kp_no = datetime.now().strftime("%Y-%m-%d/%H%M")

    client_line = ""
    if user:
        fio = (user.get("full_name") or "").strip()
        phone = (user.get("phone") or "").strip()
        reg_type = (user.get("reg_type") or "").strip()
        parts = [p for p in [fio, reg_type, phone] if p]
        client_line = " • ".join(parts)

    totals = _calc_totals(items)

    info = doc.add_table(rows=1, cols=2)
    c1, c2 = info.rows[0].cells
    _docx_set_cell_shading(c1, HEADER_SOFT)
    _docx_set_cell_shading(c2, HEADER_SOFT)

    p = c1.paragraphs[0]
    p.add_run(f"КП №: {kp_no}\n").bold = True
    p.add_run(f"Дата: {today}\n").bold = True
    p.add_run(f"Позиции: {len(items)}\n").bold = True
    if client_line:
        p.add_run(f"Клиент: {client_line}\n").bold = True

    p = c2.paragraphs[0]
    if totals.qty_total is None:
        p.add_run("Кол-во (итого): —\n").bold = True
        p.add_run("Итого со скидкой: —\n").bold = True
        p.add_run("Итого без скидки: —\n").bold = True
        p.add_run("Экономия: —\n").bold = True
        note = doc.add_paragraph("💡 Кол-во не задано — итоги не рассчитываются.")
        if note.runs:
            note.runs[0].font.size = Pt(9)
            note.runs[0].font.name = "Arial"
    else:
        p.add_run(f"Кол-во (итого): {totals.qty_total}\n").bold = True
        p.add_run(f"Итого со скидкой: {_money(totals.sum_final)}\n").bold = True
        p.add_run(f"Итого без скидки: {_money(totals.sum_price)}\n").bold = True
        p.add_run(f"Экономия: {_money(totals.economy)}\n").bold = True

    doc.add_paragraph("")

    doc.add_paragraph("Условия").runs[0].bold = True
    cond = doc.add_table(rows=3, cols=2)
    cond.cell(0, 0).text = "Срок поставки"
    cond.cell(0, 1).text = "—"
    cond.cell(1, 0).text = "Оплата"
    cond.cell(1, 1).text = "—"
    cond.cell(2, 0).text = "Доставка"
    cond.cell(2, 1).text = "—"

    doc.add_paragraph("")

    doc.add_paragraph("Состав КП").runs[0].bold = True
    tbl = doc.add_table(rows=1, cols=5)
    hdr = tbl.rows[0].cells
    hdr[0].text = "№"
    hdr[1].text = "Наименование"
    hdr[2].text = "Кол-во"
    hdr[3].text = "Цена"
    hdr[4].text = "Сумма"

    for i, it in enumerate(items, start=1):
        row = tbl.add_row().cells
        row[0].text = str(i)
        row[1].text = _short(_item_title(it), 80)
        qty = it.get("doc_qty")
        row[2].text = str(qty) if qty is not None else "—"

        final = _safe_float(it.get("final_price") if it.get("final_price") is not None else it.get("price"))
        row[3].text = _money(final) if final else "—"
        if qty is None or not final:
            row[4].text = "—"
        else:
            row[4].text = _money(final * int(qty))

    doc.save(out_path)


# -------------------- XLSX helpers --------------------

def _xlsx_item_code(it: dict[str, Any]) -> str:
    code = (
        it.get("code")
        or it.get("article")
        or it.get("sku")
        or it.get("vendor_code")
        or (it.get("product_code") if isinstance(it.get("product_code"), str) else None)
        or ""
    )
    return str(code).strip()


def _xlsx_price_value(it: dict[str, Any]) -> float | None:
    raw = it.get("final_price") if it.get("final_price") not in (None, "") else it.get("price")
    if raw in (None, ""):
        return None
    try:
        return float(raw)
    except Exception:
        return None


def _xlsx_qty_value(it: dict[str, Any]) -> int | None:
    raw = it.get("doc_qty")
    if raw in (None, ""):
        return None
    try:
        return int(raw)
    except Exception:
        return None


def _prepare_image_for_xlsx(path: str | None) -> str:
    img_path = (path or "").strip()
    if not img_path or not os.path.exists(img_path):
        img_path = default_image_path()

    if Image is None or ImageOps is None:
        return img_path

    try:
        src = Path(img_path)
        if not src.exists() or not src.is_file():
            return default_image_path()

        cache_dir = Path(_kp_img_dir()) / "xlsx"
        cache_dir.mkdir(parents=True, exist_ok=True)

        stat = src.stat()
        out_path = cache_dir / f"{src.stem}_{int(stat.st_mtime)}_{stat.st_size}_xlsx.jpg"

        if out_path.exists() and out_path.stat().st_size > 0:
            return str(out_path)

        with Image.open(src) as img:
            img = ImageOps.exif_transpose(img)

            if img.mode in ("RGBA", "LA", "P"):
                bg = Image.new("RGB", img.size, "white")
                if img.mode == "RGBA":
                    bg.paste(img, mask=img.getchannel("A"))
                else:
                    rgba = img.convert("RGBA")
                    bg.paste(rgba, mask=rgba.getchannel("A"))
                img = bg
            else:
                img = img.convert("RGB")

            resampling = getattr(getattr(Image, "Resampling", Image), "LANCZOS", 1)
            img.thumbnail((260, 260), resampling)
            img.save(out_path, "JPEG", quality=84, optimize=True)

        return str(out_path) if out_path.exists() else img_path
    except Exception:
        return default_image_path()


def _style_xlsx_sheet(ws) -> None:
    header_fill = PatternFill("solid", fgColor=HEADER_SOFT.replace("#", ""))
    header_font = Font(bold=True, color=TEXT_DARK.replace("#", ""))
    thin = Side(style="thin", color=BORDER_SOFT.replace("#", ""))
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for cell in ws[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=True)

    widths = {
        "A": 16,
        "B": 18,
        "C": 38,
        "D": 55,
        "E": 12,
        "F": 16,
        "G": 16,
        "H": 16,
        "I": 38,
        "J": 32,
        "K": 32,
    }
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    ws.freeze_panes = "B2"
    ws.auto_filter.ref = ws.dimensions


def create_kp_xlsx(
    *,
    out_path: str,
    user: dict[str, Any] | None,
    items: list[dict[str, Any]],
    logo_path: str | None,
) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "КП"

    headers = [
        "Фото",
        "Артикул",
        "Название",
        "Описание",
        "Кол-во",
        "Цена из прайса",
        "Ручная цена",
        "Сумма",
        "Ссылка",
        "Комментарий менеджера",
        "Примечание",
    ]
    ws.append(headers)
    ws.row_dimensions[1].height = 30

    for idx, it in enumerate(items, start=2):
        qty = _xlsx_qty_value(it)
        price = _xlsx_price_value(it)
        title = _item_title(it)
        desc = (it.get("description") or "").strip()
        url = (it.get("url") or "").strip()

        ws.cell(row=idx, column=2, value=_xlsx_item_code(it))
        ws.cell(row=idx, column=3, value=title)
        ws.cell(row=idx, column=4, value=desc)
        ws.cell(row=idx, column=5, value=qty)
        ws.cell(row=idx, column=6, value=price)
        ws.cell(row=idx, column=7, value=None)
        ws.cell(
            row=idx,
            column=8,
            value=f'=IF(OR(E{idx}="",IF(G{idx}<>"",G{idx},F{idx})=""),"",E{idx}*IF(G{idx}<>"",G{idx},F{idx}))',
        )
        ws.cell(row=idx, column=9, value=url)
        ws.cell(row=idx, column=10, value="")
        ws.cell(row=idx, column=11, value="")

        for col in (6, 7, 8):
            ws.cell(row=idx, column=col).number_format = '# ##0.00'

        img_path = (it.get("image_path") or "").strip()
        if not img_path or not os.path.exists(img_path):
            img_path = default_image_path()
        img_path = _prepare_image_for_xlsx(img_path)

        try:
            ximg = XLImage(img_path)
            ximg.width = 92
            ximg.height = 72
            ws.add_image(ximg, f"A{idx}")
        except Exception:
            pass

        ws.row_dimensions[idx].height = 62

    total_row = len(items) + 2
    ws.cell(row=total_row, column=3, value="ИТОГО")
    ws.cell(row=total_row, column=3).font = Font(bold=True)
    ws.cell(row=total_row, column=5, value=f"=SUM(E2:E{total_row - 1})")
    ws.cell(row=total_row, column=8, value=f"=SUM(H2:H{total_row - 1})")
    ws.cell(row=total_row, column=8).number_format = '# ##0.00'

    _style_xlsx_sheet(ws)

    for col_idx in range(1, len(headers) + 1):
        ws.cell(row=total_row, column=col_idx).border = ws.cell(row=1, column=col_idx).border
        ws.cell(row=total_row, column=col_idx).alignment = Alignment(vertical="center", wrap_text=True)

    wb.save(out_path)


# -------------------- Wizard UI --------------------

def _mode_kb() -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text="⚡ Сгенерировать (кол-во — «—»)", callback_data="kp_build:mode:dash"),
    )
    kb.row(
        InlineKeyboardButton(text="✍️ Указать кол-во + (опц.) ссылка", callback_data="kp_build:mode:manual"),
    )
    kb.row(
        InlineKeyboardButton(text="❌ Отмена", callback_data="kp_build:cancel"),
    )
    return kb


def _qty_kb() -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text="⏭ Пропустить (—)", callback_data="kp_build:qty:skip"),
        InlineKeyboardButton(text="🔗 Добавить ссылку", callback_data="kp_build:url:add"),
    )
    kb.row(
        InlineKeyboardButton(text="⬅️ Назад", callback_data="kp_build:qty:back"),
    )
    kb.row(
        InlineKeyboardButton(text="❌ Отмена", callback_data="kp_build:cancel"),
    )
    return kb


def _url_kb() -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    kb.row(
        InlineKeyboardButton(text="⏭ Пропустить ссылку", callback_data="kp_build:url:skip"),
        InlineKeyboardButton(text="⬅️ Назад", callback_data="kp_build:url:back"),
    )
    kb.row(
        InlineKeyboardButton(text="❌ Отмена", callback_data="kp_build:cancel"),
    )
    return kb


def _missing_indices(items: list[dict[str, Any]]) -> list[int]:
    """
    Возвращает индексы позиций, где не заполнено кол-во ИЛИ нет ссылки.
    """
    out: list[int] = []
    for i, it in enumerate(items):
        qty_missing = it.get("doc_qty") is None
        url_missing = not (it.get("url") or "").strip()
        if qty_missing or url_missing:
            out.append(i)
    return out


def _render_review(items: list[dict[str, Any]]) -> str:
    qty_missing = [i for i, it in enumerate(items) if it.get("doc_qty") is None]
    url_missing = [i for i, it in enumerate(items) if not (it.get("url") or "").strip()]

    lines = []
    lines.append("🧾 <b>Проверка перед формированием КП</b>\n")
    lines.append(f"Позиции всего: <b>{len(items)}</b>")
    lines.append(f"Без кол-ва: <b>{len(qty_missing)}</b>")
    lines.append(f"Без ссылки: <b>{len(url_missing)}</b>\n")

    miss = _missing_indices(items)
    if not miss:
        lines.append("✅ Всё заполнено. Можно формировать КП.")
        return "\n".join(lines)

    lines.append("Незаполненные позиции (первые 15):")
    shown = 0
    for idx in miss:
        it = items[idx]
        title = _short(_item_title(it), 70)
        qm = "нет кол-ва" if it.get("doc_qty") is None else None
        um = "нет ссылки" if not (it.get("url") or "").strip() else None
        parts = [p for p in (qm, um) if p]
        lines.append(f"{idx + 1}) {title} — {', '.join(parts)}")
        shown += 1
        if shown >= 15:
            break

    if len(miss) > shown:
        lines.append(f"… и ещё {len(miss) - shown} поз.")

    return "\n".join(lines)


def _review_kb(has_missing: bool) -> InlineKeyboardBuilder:
    kb = InlineKeyboardBuilder()
    if has_missing:
        kb.row(InlineKeyboardButton(text="✍️ Заполнить пропуски", callback_data="kp_build:review:fill"))
    kb.row(InlineKeyboardButton(text="✅ Сформировать КП", callback_data="kp_build:review:gen"))
    kb.row(InlineKeyboardButton(text="❌ Отмена", callback_data="kp_build:cancel"))
    return kb


async def _show_review(call: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    await state.set_state(KpBuildForm.review)
    text = _render_review(items)
    has_missing = bool(_missing_indices(items))
    await safe_edit_text(call.message, text, reply_markup=_review_kb(has_missing).as_markup())


def _render_qty_step(items: list[dict[str, Any]], idx: int) -> str:
    it = items[idx]
    title = _short(_item_title(it), 160)
    stock = _stock_of_item(it)
    stock_line = "Наличие: <b>—</b>" if stock is None else (f"Наличие: <b>{stock} шт</b>" if stock > 0 else "Наличие: <b>0 шт</b> ❌")

    url = (it.get("url") or "").strip()
    url_line = f"🔗 Ссылка: <b>добавлена</b>" if url else "🔗 Ссылка: <b>нет</b>"

    return (
        "🧾 <b>Формирование КП</b>\n\n"
        f"Позиция <b>{idx + 1}</b> / <b>{len(items)}</b>\n"
        f"📦 <b>{title}</b>\n"
        f"{stock_line}\n"
        f"{url_line}\n\n"
        "✍️ Отправьте <b>число</b> (кол-во) сообщением.\n"
        "Или используйте кнопки ниже (можно добавить ссылку, чтобы подтянуть фото/описание)."
    )


def _render_url_step(items: list[dict[str, Any]], idx: int) -> str:
    it = items[idx]
    title = _short(_item_title(it), 160)
    return (
        "🔗 <b>Добавление ссылки</b>\n\n"
        f"Позиция <b>{idx + 1}</b> / <b>{len(items)}</b>\n"
        f"📦 <b>{title}</b>\n\n"
        "Отправьте ссылку на страницу товара (http/https).\n"
        "Бот попробует подтянуть <b>фото</b> и <b>недостающую информацию</b> для КП."
    )


async def _apply_enrich_for_item(it: dict[str, Any], url: str, uid: int, idx: int) -> bool:
    """
    Заполняем ТОЛЬКО недостающее по данным страницы + скачиваем фото, чтобы попало в PDF.
    """
    try:
        info = await enrich_from_url(url)
    except Exception:
        return False

    if not isinstance(info, dict):
        return False

    it["url"] = url

    title = (info.get("title") or "").strip() if isinstance(info.get("title"), str) else ""
    desc = (info.get("description") or "").strip() if isinstance(info.get("description"), str) else ""
    img_url = ""
    if isinstance(info.get("image_url"), str):
        img_url = (info.get("image_url") or "").strip()
    elif isinstance(info.get("image"), str):
        img_url = (info.get("image") or "").strip()

    if _is_empty(it.get("title")) and title:
        it["title"] = title
    if _is_empty(it.get("description")) and desc:
        it["description"] = desc
    if _is_empty(it.get("image_url")) and img_url:
        it["image_url"] = img_url

    # если нет локальной картинки — пробуем скачать
    await _ensure_item_image_downloaded(it, uid=uid, idx=idx)
    return True


async def _prepare_items_for_docs(uid: int, items: list[dict[str, Any]]) -> None:
    """
    На всякий случай: если у позиции есть image_url, но нет image_path — скачиваем перед генерацией.
    """
    for i, it in enumerate(items, start=1):
        await _ensure_item_image_downloaded(it, uid=uid, idx=i)


# -------------------- Entry --------------------

@router.message(IsApprovedUser(), F.text == "🧾 Сформировать КП")
async def build_kp_from_message(message: Message, settings: Settings, state: FSMContext) -> None:
    await _start_build_wizard(message, settings, state)


@router.callback_query(IsApprovedUser(), UserKpCb.filter(F.action == "build"))
async def build_kp_from_callback(call: CallbackQuery, callback_data: UserKpCb, settings: Settings, state: FSMContext) -> None:
    await _start_build_wizard(call, settings, state)


async def _start_build_wizard(target: Message | CallbackQuery, settings: Settings, state: FSMContext) -> None:
    tg_id = target.from_user.id

    total = await count_kp_items(settings.db_path, tg_id)
    if total <= 0:
        if isinstance(target, CallbackQuery):
            await target.answer("КП пустое. Добавьте товары.", show_alert=True)
        else:
            await target.answer("КП пустое. Добавьте товары.")
        return

    user = await get_user(settings.db_path, tg_id)
    items = await list_kp_items(settings.db_path, tg_id, limit=total, offset=0)

    await _enrich_items_with_stock(settings, items)

    # ✅ КП может быть от разных поставщиков
    supplier_ids = sorted({int(it.get("supplier_id")) for it in items if it.get("supplier_id")})
    primary_supplier_id = supplier_ids[0] if supplier_ids else None  # для совместимости, если где-то дальше ожидается одно число

    await state.clear()
    await state.set_state(KpBuildForm.choosing_mode)
    await state.update_data(
        kp_user=user,
        kp_items=items,
        supplier_id=primary_supplier_id,
        supplier_ids=supplier_ids,
        kp_uid=tg_id,
    )

    text = (
        "🧾 <b>Формирование КП</b>\n\n"
        "Как заполнить поле <b>Кол-во</b>?\n\n"
        "1) «—» — быстро, без ввода количества (итоги не считаются)\n"
        "2) Ввести — бот попросит кол-во по каждой позиции и проверит наличие.\n\n"
        "💡 В режиме 2 также можно <b>добавить ссылку</b> на товар, чтобы подтянуть фото/описание в КП."
    )

    kb = _mode_kb().as_markup()

    if isinstance(target, CallbackQuery):
        try:
            await target.message.delete()
        except Exception:
            pass
        msg = await target.message.bot.send_message(target.message.chat.id, text, reply_markup=kb)
        await target.answer()
    else:
        msg = await target.answer(text, reply_markup=kb)

    await state.update_data(wizard_msg_id=msg.message_id)


# -------------------- Mode callbacks --------------------

@router.callback_query(IsApprovedUser(), F.data == "kp_build:cancel")
async def kp_build_cancel(call: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    mid = int(data.get("wizard_msg_id") or 0)
    await state.clear()
    await safe_delete_message(call.message.bot, call.message.chat.id, mid)
    await call.answer("Отменено ✅", show_alert=False)


@router.callback_query(IsApprovedUser(), F.data == "kp_build:mode:dash")
async def kp_build_mode_dash(call: CallbackQuery, settings: Settings, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    user = data.get("kp_user")
    uid = int(data.get("kp_uid") or call.from_user.id)

    for it in items:
        it["doc_qty"] = None

    # на всякий случай подтянем картинки по image_url (если они уже есть)
    await safe_edit_text(call.message, "⏳ Готовлю данные (картинки)…", reply_markup=None)
    await _prepare_items_for_docs(uid, items)

    await _generate_and_send(call, settings, user, items, state)


@router.callback_query(IsApprovedUser(), F.data == "kp_build:mode:manual")
async def kp_build_mode_manual(call: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    if not items:
        await call.answer("КП пустое.", show_alert=True)
        return

    await state.set_state(KpBuildForm.waiting_qty)
    await state.update_data(qty_index=0)

    text = _render_qty_step(items, 0)
    await safe_edit_text(call.message, text, reply_markup=_qty_kb().as_markup())
    await call.answer()


# -------------------- Qty navigation callbacks --------------------

@router.callback_query(IsApprovedUser(), KpBuildForm.waiting_qty, F.data == "kp_build:qty:back")
async def kp_qty_back(call: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    idx = int(data.get("qty_index") or 0)

    idx = max(0, idx - 1)
    await state.update_data(qty_index=idx)

    text = _render_qty_step(items, idx)
    await safe_edit_text(call.message, text, reply_markup=_qty_kb().as_markup())
    await call.answer()


@router.callback_query(IsApprovedUser(), KpBuildForm.waiting_qty, F.data == "kp_build:qty:skip")
async def kp_qty_skip(call: CallbackQuery, settings: Settings, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    user = data.get("kp_user")
    uid = int(data.get("kp_uid") or call.from_user.id)
    idx = int(data.get("qty_index") or 0)

    items[idx]["doc_qty"] = None

    idx += 1
    if idx >= len(items):
        await state.update_data(kp_items=items, review_editing=False)
        await _show_review(call, state)
        return

    await state.update_data(qty_index=idx, kp_items=items)
    text = _render_qty_step(items, idx)
    await safe_edit_text(call.message, text, reply_markup=_qty_kb().as_markup())
    await call.answer()


# -------------------- URL actions (callbacks) --------------------

@router.callback_query(IsApprovedUser(), KpBuildForm.waiting_qty, F.data == "kp_build:url:add")
async def kp_url_add_from_qty(call: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    idx = int(data.get("qty_index") or 0)

    await state.set_state(KpBuildForm.waiting_url)

    text = _render_url_step(items, idx)
    await safe_edit_text(call.message, text, reply_markup=_url_kb().as_markup())
    await call.answer()


@router.callback_query(IsApprovedUser(), KpBuildForm.waiting_url, F.data == "kp_build:url:back")
async def kp_url_back(call: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    idx = int(data.get("qty_index") or 0)

    await state.set_state(KpBuildForm.waiting_qty)

    text = _render_qty_step(items, idx)
    await safe_edit_text(call.message, text, reply_markup=_qty_kb().as_markup())
    await call.answer()


@router.callback_query(IsApprovedUser(), KpBuildForm.waiting_url, F.data == "kp_build:url:skip")
async def kp_url_skip(call: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    idx = int(data.get("qty_index") or 0)

    # ничего не меняем, просто назад к количеству
    await state.set_state(KpBuildForm.waiting_qty)
    text = _render_qty_step(items, idx)
    await safe_edit_text(call.message, text, reply_markup=_qty_kb().as_markup())
    await call.answer()


# -------------------- URL input (messages) --------------------

@router.message(IsApprovedUser(), KpBuildForm.waiting_url, F.text)
async def kp_url_input(message: Message, settings: Settings, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    idx = int(data.get("qty_index") or 0)
    uid = int(data.get("kp_uid") or message.from_user.id)

    txt = (message.text or "").strip()
    await safe_delete_message(message.bot, message.chat.id, message.message_id)

    if not items or idx < 0 or idx >= len(items):
        await state.clear()
        return

    if not _is_valid_url(txt):
        err = await message.bot.send_message(message.chat.id, "⚠️ Отправьте корректную ссылку (http/https) или нажмите «Пропустить ссылку».")
        await asyncio.sleep(1.4)
        await safe_delete_message(message.bot, message.chat.id, err.message_id)
        return

    # покажем мини-прогресс (без засорения)
    progress = await message.bot.send_message(message.chat.id, "⏳ Читаю страницу и подтягиваю фото/описание…")

    ok = await _apply_enrich_for_item(items[idx], txt, uid=uid, idx=idx + 1)

    await safe_delete_message(message.bot, message.chat.id, progress.message_id)

    if not ok:
        err = await message.bot.send_message(message.chat.id, "❌ Не получилось прочитать страницу. Попробуйте другую ссылку или нажмите «Пропустить ссылку».")
        await asyncio.sleep(1.6)
        await safe_delete_message(message.bot, message.chat.id, err.message_id)
        return

    await state.update_data(kp_items=items)

    if bool(data.get("review_editing")):
        # здесь Message-хендлер, поэтому обновляем wizard msg вручную
        wizard_mid = int(data.get("wizard_msg_id") or 0)
        if wizard_mid:
            try:
                await message.bot.edit_message_text(
                    chat_id=message.chat.id,
                    message_id=wizard_mid,
                    text=_render_review(items),
                    reply_markup=_review_kb(bool(_missing_indices(items))).as_markup(),
                )
                await state.set_state(KpBuildForm.review)
            except Exception:
                m = await message.bot.send_message(
                    message.chat.id,
                    _render_review(items),
                    reply_markup=_review_kb(bool(_missing_indices(items))).as_markup(),
                )
                await state.update_data(wizard_msg_id=m.message_id)
                await state.set_state(KpBuildForm.review)
        return

    # обычный сценарий: возвращаемся к вводу количества
    await state.set_state(KpBuildForm.waiting_qty)

    wizard_mid = int(data.get("wizard_msg_id") or 0)
    if wizard_mid:
        try:
            await message.bot.edit_message_text(
                chat_id=message.chat.id,
                message_id=wizard_mid,
                text=_render_qty_step(items, idx),
                reply_markup=_qty_kb().as_markup(),
            )
        except Exception:
            m = await message.bot.send_message(message.chat.id, _render_qty_step(items, idx), reply_markup=_qty_kb().as_markup())
            await state.update_data(wizard_msg_id=m.message_id)


# -------------------- Qty input (messages) --------------------

@router.message(IsApprovedUser(), KpBuildForm.waiting_qty, F.text)
async def kp_qty_input(message: Message, settings: Settings, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    user = data.get("kp_user")
    uid = int(data.get("kp_uid") or message.from_user.id)
    wizard_mid = int(data.get("wizard_msg_id") or 0)
    idx = int(data.get("qty_index") or 0)

    txt = (message.text or "").strip()
    await safe_delete_message(message.bot, message.chat.id, message.message_id)

    if not items or idx < 0 or idx >= len(items):
        await state.clear()
        return

    if not re.fullmatch(r"\d+", txt):
        err = await message.bot.send_message(message.chat.id, "⚠️ Введите <b>целое число</b> (например: 6).")
        await asyncio.sleep(1.2)
        await safe_delete_message(message.bot, message.chat.id, err.message_id)
        return

    qty = int(txt)
    if qty <= 0:
        err = await message.bot.send_message(message.chat.id, "⚠️ Кол-во должно быть больше 0.")
        await asyncio.sleep(1.2)
        await safe_delete_message(message.bot, message.chat.id, err.message_id)
        return

    stock = _stock_of_item(items[idx])
    if stock is not None and qty > stock:
        err = await message.bot.send_message(
            message.chat.id,
            f"❌ Недостаточно на складе.\nДоступно: <b>{stock} шт</b>\nВведите другое кол-во.",
        )
        await asyncio.sleep(1.6)
        await safe_delete_message(message.bot, message.chat.id, err.message_id)
        return

    items[idx]["doc_qty"] = qty

    if bool(data.get("review_editing")):
        await state.update_data(kp_items=items)
        try:
            await message.bot.edit_message_text(
                chat_id=message.chat.id,
                message_id=wizard_mid,
                text=_render_review(items),
                reply_markup=_review_kb(bool(_missing_indices(items))).as_markup(),
            )
            await state.set_state(KpBuildForm.review)
        except Exception:
            m = await message.bot.send_message(
                message.chat.id,
                _render_review(items),
                reply_markup=_review_kb(bool(_missing_indices(items))).as_markup(),
            )
            await state.update_data(wizard_msg_id=m.message_id)
            await state.set_state(KpBuildForm.review)
        return

    next_idx = idx + 1

    # если это был последний товар — показываем review
    if next_idx >= len(items):
        await state.update_data(kp_items=items, review_editing=False)

        text = _render_review(items)
        has_missing = bool(_missing_indices(items))

        if wizard_mid:
            try:
                await message.bot.edit_message_text(
                    chat_id=message.chat.id,
                    message_id=wizard_mid,
                    text=text,
                    reply_markup=_review_kb(has_missing).as_markup(),
                )
            except Exception:
                m = await message.bot.send_message(
                    message.chat.id,
                    text,
                    reply_markup=_review_kb(has_missing).as_markup(),
                )
                await state.update_data(wizard_msg_id=m.message_id)
        else:
            m = await message.bot.send_message(
                message.chat.id,
                text,
                reply_markup=_review_kb(has_missing).as_markup(),
            )
            await state.update_data(wizard_msg_id=m.message_id)

        await state.set_state(KpBuildForm.review)
        return

    # иначе идём на следующий товар
    await state.update_data(qty_index=next_idx, kp_items=items)

    if wizard_mid:
        try:
            await message.bot.edit_message_text(
                chat_id=message.chat.id,
                message_id=wizard_mid,
                text=_render_qty_step(items, next_idx),
                reply_markup=_qty_kb().as_markup(),
            )
        except Exception:
            m = await message.bot.send_message(
                message.chat.id,
                _render_qty_step(items, next_idx),
                reply_markup=_qty_kb().as_markup(),
            )
            await state.update_data(wizard_msg_id=m.message_id)
    else:
        m = await message.bot.send_message(
            message.chat.id,
            _render_qty_step(items, next_idx),
            reply_markup=_qty_kb().as_markup(),
        )
        await state.update_data(wizard_msg_id=m.message_id)


# -------------------- Generators --------------------

async def _send_document_safely(
    bot,
    chat_id: int,
    path: str,
    *,
    caption: str | None = None,
) -> bool:
    if not path or not os.path.exists(path):
        return False

    if os.path.getsize(path) > TELEGRAM_DOCUMENT_LIMIT_BYTES:
        return False

    try:
        await bot.send_document(chat_id, FSInputFile(path), caption=caption)
        return True
    except TelegramEntityTooLarge:
        return False
    except Exception:
        log.exception("Failed to send generated document: %s", path)
        return False


async def _send_kp_files(bot, chat_id: int, pdf_path: str, xlsx_path: str) -> None:
    pdf_mb = _file_size_mb(pdf_path)
    xlsx_mb = _file_size_mb(xlsx_path)

    pdf_sent = await _send_document_safely(
        bot,
        chat_id,
        pdf_path,
        caption="✅ КП сформировано. Отправляю PDF и Excel.",
    )

    if pdf_sent:
        xlsx_sent = await _send_document_safely(bot, chat_id, xlsx_path)
        if not xlsx_sent:
            await bot.send_message(
                chat_id,
                f"⚠️ PDF отправлен, но Excel-файл получился слишком большим для Telegram.\n"
                f"Размер Excel: <b>{xlsx_mb} МБ</b>.",
            )
        return

    xlsx_sent = await _send_document_safely(
        bot,
        chat_id,
        xlsx_path,
        caption=(
            "⚠️ PDF получился слишком большим для Telegram.\n"
            "Отправляю Excel-версию КП."
        ),
    )

    if xlsx_sent:
        await bot.send_message(
            chat_id,
            f"ℹ️ PDF не отправлен из-за размера.\n"
            f"Размер PDF: <b>{pdf_mb} МБ</b>.\n\n"
            "КП отправлено в Excel-формате.",
        )
        return

    await bot.send_message(
        chat_id,
        "❌ КП сформировано, но файлы получились слишком большими для Telegram.\n\n"
        f"PDF: <b>{pdf_mb} МБ</b>\n"
        f"Excel: <b>{xlsx_mb} МБ</b>\n\n"
        "Уменьшите количество позиций или изображения товаров."
    )

async def _generate_and_send_from_message(
    message: Message,
    settings: Settings,
    user: dict[str, Any] | None,
    items: list[dict[str, Any]],
    state: FSMContext,
    wizard_mid: int,
    progress_mid: int,
    uid: int,
) -> None:
    logo_path = _find_logo_path()
    base = _build_base_filename()
    pdf_path = os.path.join(_kp_dir(), f"{base}.pdf")
    xlsx_path = os.path.join(_kp_dir(), f"{base}.xlsx")

    try:
        await _prepare_items_for_docs(uid, items)
        create_kp_pdf(out_path=pdf_path, user=user, items=items, logo_path=logo_path)
        create_kp_xlsx(out_path=xlsx_path, user=user, items=items, logo_path=logo_path)
    except Exception:
        log.exception("KP build failed")
        await safe_delete_message(message.bot, message.chat.id, progress_mid)
        await message.bot.send_message(
            message.chat.id,
            "❌ Не получилось сформировать КП.\n"
            "Проверьте, что лого лежит в media/logo/png.\n"
            "Для PDF используются системные шрифты (DejaVu/Arial).",
        )
        await state.clear()
        return

    await safe_delete_message(message.bot, message.chat.id, progress_mid)
    await safe_delete_message(message.bot, message.chat.id, wizard_mid)

    await _send_kp_files(message.bot, message.chat.id, pdf_path, xlsx_path)

    await state.clear()


async def _generate_and_send(
    call: CallbackQuery,
    settings: Settings,
    user: dict[str, Any] | None,
    items: list[dict[str, Any]],
    state: FSMContext,
) -> None:
    logo_path = _find_logo_path()
    base = _build_base_filename()
    pdf_path = os.path.join(_kp_dir(), f"{base}.pdf")
    xlsx_path = os.path.join(_kp_dir(), f"{base}.xlsx")

    await safe_edit_text(call.message, "⏳ Формирую PDF и Excel…", reply_markup=None)

    uid = call.from_user.id

    try:
        await _prepare_items_for_docs(uid, items)
        create_kp_pdf(out_path=pdf_path, user=user, items=items, logo_path=logo_path)
        create_kp_xlsx(out_path=xlsx_path, user=user, items=items, logo_path=logo_path)
    except Exception:
        log.exception("KP build failed")
        await safe_edit_text(
            call.message,
            "❌ Не получилось сформировать КП.\n"
            "Проверьте, что лого лежит в media/logo/png.\n"
            "Для PDF используются системные шрифты (DejaVu/Arial).",
            reply_markup=None,
        )
        await state.clear()
        await call.answer("Ошибка формирования", show_alert=True)
        return

    data = await state.get_data()
    mid = int(data.get("wizard_msg_id") or 0)
    await safe_delete_message(call.message.bot, call.message.chat.id, mid)

    await _send_kp_files(call.message.bot, call.message.chat.id, pdf_path, xlsx_path)
    await call.answer()

    await state.clear()


@router.callback_query(IsApprovedUser(), KpBuildForm.review, F.data == "kp_build:review:fill")
async def kp_review_fill(call: CallbackQuery, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    miss = _missing_indices(items)
    if not miss:
        await call.answer("Пропусков нет.")
        await _show_review(call, state)
        return

    idx = miss[0]
    await state.update_data(qty_index=idx, kp_items=items, review_editing=True)

    # если нет кол-ва — идём на qty, иначе на url
    if items[idx].get("doc_qty") is None:
        await state.set_state(KpBuildForm.waiting_qty)
        await safe_edit_text(call.message, _render_qty_step(items, idx), reply_markup=_qty_kb().as_markup())
    else:
        await state.set_state(KpBuildForm.waiting_url)
        await safe_edit_text(call.message, _render_url_step(items, idx), reply_markup=_url_kb().as_markup())

    await call.answer()


@router.callback_query(IsApprovedUser(), KpBuildForm.review, F.data == "kp_build:review:gen")
async def kp_review_generate(call: CallbackQuery, settings: Settings, state: FSMContext) -> None:
    data = await state.get_data()
    items = list(data.get("kp_items") or [])
    user = data.get("kp_user")

    # финальная генерация
    await state.update_data(review_editing=False)
    await _generate_and_send(call, settings, user, items, state)